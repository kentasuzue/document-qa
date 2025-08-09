# Copyright (C) 2025 Kenta Suzue

import streamlit as st
from langchain_deeplake.vectorstores import DeeplakeVectorStore
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document
import docx
# import numpy as np
import openai
import os
from PyPDF2 import PdfReader
import tempfile
import spacy
from deeplake import delete

MAX_RESUMES_SUMMARIZED = 10
DATASET_PATH = "hub://kentasuzue/resume-matcher"

# clear dataset
if "dataset_deleted" not in st.session_state:
    try:
        delete(DATASET_PATH)
        st.session_state.dataset_deleted = True
        # st.info("✅ Vector store cleared.")
    except Exception as e:
        st.warning(f"Warning: Could not delete dataset: {e}")

# Load API keys
openai.api_key = st.secrets["OPENAI_API_KEY"]
os.environ["ACTIVELOOP_TOKEN"] = st.secrets["ACTIVELOOP_TOKEN"]

# Load spaCy English model once globally
nlp = spacy.load("en_core_web_sm")

# Commit job description text
def commit_job_description():
    st.session_state.job_text = st.session_state.job_description_input
    
# File Read Functions
def parse_file(file):
    if file.name.endswith(".pdf"):
        reader = PdfReader(file)
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")
    else:
        return "Unsupported file type"

# Get OpenAI embedding
@st.cache_resource(show_spinner=False)
def get_embedding():
    return OpenAIEmbeddings(model="text-embedding-3-large")

def clean_text(text):
    """Remove non-text characters and ensure UTF-8 compliance."""
    if isinstance(text, bytes):
        text = text.decode("utf-8", errors="ignore")
    return ''.join(char for char in text if char.isprintable())
    
def extract_candidate_name(resume_text):
    """
    Attempts to extract the candidate's name from the resume.
    Assumes the name is within the first few lines.
    """
    lines = resume_text.strip().splitlines()
    for line in lines[:5]:  # Check the top 5 lines
        line = line.strip()
        if line and len(line.split()) <= 5 and any(char.isalpha() for char in line):
            return line
    return "Unknown"

def extract_candidate_name_fancy(resume_text):
    """
    Attempts to extract the candidate's name from the resume using spaCy NER.
    If none found, fallback to heuristic.
    """
    lines = resume_text.strip().splitlines()
    first_lines_text = "\n".join(lines[:10])
    doc = nlp(first_lines_text)
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    # fallback heuristic: pick first short line with letters
    for line in lines[:10]:
        line = line.strip()
        if line and len(line.split()) <= 5 and any(char.isalpha() for char in line):
            return line
    return "Unknown"

# Generate GPT summary
@st.cache_data(show_spinner=False)
def generate_summary_with_gpt(job_text, resume_text):
    prompt = f"""
You are a helpful recruiter assistant. Given the job description and the candidate's resume, write a concise 2–3 sentence summary explaining how well the candidate fits the job. Highlight relevant experience and standout skills.

Job Description:
{job_text}

Resume:
{resume_text}

Summary:
"""
    response = openai.chat.completions.create(
        model="gpt-5-nano",
        messages=[{"role": "user", "content": prompt}],
        verbosity="high"
    )
    return response.choices[0].message.content.strip()
    
# --- App Layout ---
st.title("Candidate Recommendation Engine")
st.markdown("## New and Improved with ChatGPT-5!")
st.markdown("""
            <ul>
              <li>(1) Provide a job description.</li>
              <li>(2) Provide candidate resumes.</li>
              <li>(3) Click to match resumes to the job description.</li>
              <li>We'll match and rank candidates by relevance, and provide summaries of why each candidate is a great fit for the role!</li>
            </ul>
            """, unsafe_allow_html=True)

# --- Input Job Description ---
if "job_text" not in st.session_state:
    st.session_state.job_text = ""
if "job_description_input" not in st.session_state:
    st.session_state.job_description_input = ""
    
st.markdown("#### 👷 (1) Job Description")
# The main text area user edits:
st.text_area(
    "Paste Job Description, then type CTRL+ENTER or click '🖱️ Set Job Description'",
    height=200,
    key="job_description_input",
    on_change=commit_job_description  # triggers on Ctrl+Enter
)

# Button to commit the current input to job_text on mouse click
if st.button("🖱️ Set Job Description"):
    commit_job_description()

# Show the confirmed job_text
if st.session_state.job_text:
    st.markdown("### 📋 Confirmed Job Description:")
    st.write(st.session_state.job_text)
    
# --- Input Resumes ---
st.markdown("#### 👥 (2) Candidate Resumes")
resume_files = st.file_uploader("Upload Candidate Resumes", type=["txt", "pdf", "docx"], accept_multiple_files=True)

# --- Initialize state ---
if "pasted_resumes" not in st.session_state:
    st.session_state.pasted_resumes = []

if "resume_input" not in st.session_state:
    st.session_state.resume_input = ""

if "reset_resume_box" not in st.session_state:
    st.session_state.reset_resume_box = False

st.markdown("Or Paste in Resumes (one at a time)")

# --- Handle dynamic key for reset ---
resume_box_key = "resume_input"
if st.session_state.reset_resume_box:
    resume_box_key = f"resume_input_{len(st.session_state.pasted_resumes)}"
    st.session_state.reset_resume_box = False

# --- Text area ---
st.text_area(
    "Paste resume here",
    height=250,
    key=resume_box_key,
)

# --- Add/Clear Buttons ---
col1, col2 = st.columns([1, 1])
with col1:
    if st.button("➕ Add Resume"):
        resume_text = st.session_state.get("resume_input", "").strip()
        if resume_text:
            # Use your preferred name extraction here
            from your_module import extract_candidate_name_fancy  # if defined externally
            candidate_name = extract_candidate_name_fancy(resume_text)

            st.session_state.pasted_resumes.append(
                Document(
                    page_content=resume_text,
                    metadata={
                        "id": f"pasted_resume_{len(st.session_state.pasted_resumes) + 1}",
                        "candidate_name": candidate_name
                    }
                )
            )
            st.success(f"Resume #{len(st.session_state.pasted_resumes)} added.")
            st.session_state.resume_input = ""
            st.session_state.reset_resume_box = True
            st.rerun()
        else:
            st.warning("Paste a resume before clicking 'Add Resume'.")

with col2:
    if st.button("🧹 Clear All Pasted Resumes"):
        st.session_state.pasted_resumes = []
        st.success("Cleared all pasted resumes.")
        st.rerun()

# --- Show Pasted Resumes ---
if st.session_state.pasted_resumes:
    st.markdown("### 📄 Pasted Resumes Added")
    for i, resume_doc in enumerate(st.session_state.pasted_resumes):
        with st.expander(f"Resume #{i+1}: {resume_doc.metadata.get('candidate_name', 'Unknown')}"):
            st.text_area("Resume Text", resume_doc.page_content, height=200, key=f"resume_view_{i}")
            if st.button(f"❌ Remove Resume #{i+1}", key=f"remove_{i}"):
                st.session_state.pasted_resumes.pop(i)
                st.success(f"Removed resume #{i+1}")
                st.rerun()

# Process all resumes
resumes = []

# Load spaCy English model once
# nlp = spacy.load("en_core_web_sm")

# From uploaded files
if resume_files:
    for file in resume_files:
        text = parse_file(file)
        candidate_name = extract_candidate_name(text)
        resumes.append(Document(page_content=text, metadata={
            "id": file.name,
            "candidate_name": candidate_name
        }))

# Add pasted resumes from session state
if st.session_state.pasted_resumes:
    resumes.extend(st.session_state.pasted_resumes)

# --- Matching Trigger ---
if "run_matching" not in st.session_state:
    st.session_state.run_matching = False

if st.session_state.job_text and resumes:
    st.markdown("#### 🧠 (3) Match Resumes to Job Description")
    if st.button("🖱️ Click to Match Resumes to Job Description"):
        st.session_state.run_matching = True
        st.rerun()
        
# --- Resume Matching ---
if st.session_state.get("run_matching", False) and st.session_state.job_text and resumes:
    with st.spinner("🔄 Processing resumes and matching..."):
        embedding = get_embedding()
        
        # Create or load Deeplake vectorstore
        vs = DeeplakeVectorStore.from_documents(
            documents=resumes,
            embedding=embedding,
            dataset_path=DATASET_PATH,  # replace with your path
            overwrite=True  # overwrite dataset each run; remove if you want to append
        )

        # Search top matches for job description
        results = vs.similarity_search_with_score(st.session_state.job_text, k=min(MAX_RESUMES_SUMMARIZED, len(resumes)), distance_metric='cos')

        candidates = []
        for doc, score in results:
            summary = generate_summary_with_gpt(st.session_state.job_text, doc.page_content)
            candidates.append({
                "name": doc.metadata.get("candidate_name", doc.metadata.get("id", "Unknown")),
                "similarity": score,
                "summary": summary
            })

    # --- DISPLAY RESULTS ---
    st.header("🏆 Top Matching Candidates")
    for i, candidate in enumerate(candidates):
        st.subheader(f"#{i+1}: {candidate['name']}")
        st.write(f"**Similarity Score:** `{candidate['similarity']:.4f}`")
        st.write(f"**Summary:** {candidate['summary']}")
        st.markdown("---")
